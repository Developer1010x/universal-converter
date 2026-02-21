"""Document conversion module"""

from pathlib import Path
from typing import Optional

from . import BaseConverter, ConversionTask, ConversionResult, DependencyError


class DocumentConverter(BaseConverter):
    """Converter for document formats"""
    
    SUPPORTED_CONVERSIONS = {
        "pdf": ["docx", "txt", "html"],
        "docx": ["pdf", "txt", "html", "md"],
        "txt": ["pdf", "docx", "html", "md"],
        "html": ["pdf", "docx", "txt", "md"],
        "md": ["pdf", "docx", "txt", "html"],
    }
    PRIORITY = 15
    
    def convert(self, task: ConversionTask) -> ConversionResult:
        source = task.source_format.lower()
        target = task.target_format.lower()
        
        if source == "pdf" and target in ["docx", "txt"]:
            return self._pdf_convert(task)
        elif source == "docx" and target in ["pdf", "txt", "html"]:
            return self._docx_convert(task)
        elif source in ["txt", "md", "html"] and target in ["txt", "md", "html"]:
            return self._text_convert(task)
        
        return ConversionResult(success=False, error=f"Unsupported conversion: {source} → {target}")
    
    def _pdf_convert(self, task: ConversionTask) -> ConversionResult:
        try:
            import pypdf
            reader = pypdf.PdfReader(task.source_path)
            text = "\n".join(page.extract_text() for page in reader.pages)
            
            if task.target_format == "txt":
                task.target_path.write_text(text)
            elif task.target_format == "docx":
                return self._create_docx(text, task.target_path)
            
            return ConversionResult(success=True, output_path=str(task.target_path))
        except ImportError:
            return ConversionResult(
                success=False,
                error="pypdf required. Install: pip install universal-converter[documents]"
            )
        except Exception as e:
            return ConversionResult(success=False, error=str(e))
    
    def _docx_convert(self, task: ConversionTask) -> ConversionResult:
        try:
            from docx import Document
            doc = Document(task.source_path)
            text = "\n".join(p.text for p in doc.paragraphs)
            
            if task.target_format == "txt":
                task.target_path.write_text(text)
            elif task.target_format == "html":
                html = self._text_to_html(text)
                task.target_path.write_text(html)
            elif task.target_format == "pdf":
                return self._create_pdf(text, task.target_path)
            
            return ConversionResult(success=True, output_path=str(task.target_path))
        except ImportError:
            return ConversionResult(
                success=False,
                error="python-docx required. Install: pip install universal-converter[documents]"
            )
        except Exception as e:
            return ConversionResult(success=False, error=str(e))
    
    def _text_convert(self, task: ConversionTask) -> ConversionResult:
        try:
            content = task.source_path.read_text()
            
            if task.target_format == "html":
                content = self._text_to_html(content)
            elif task.target_format == "md":
                content = self._html_to_md(content)
            
            task.target_path.write_text(content)
            return ConversionResult(success=True, output_path=str(task.target_path))
        except Exception as e:
            return ConversionResult(success=False, error=str(e))
    
    def _text_to_html(self, text: str) -> str:
        import html
        lines = text.split("\n")
        html_lines = ['<html><body>']
        for line in lines:
            if line.strip():
                html_lines.append(f"<p>{html.escape(line)}</p>")
        html_lines.append('</body></html>')
        return "\n".join(html_lines)
    
    def _html_to_md(self, html: str) -> str:
        import re
        md = html
        md = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1\n', md, flags=re.DOTALL)
        md = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1\n', md, flags=re.DOTALL)
        md = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1\n', md, flags=re.DOTALL)
        md = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n', md, flags=re.DOTALL)
        md = re.sub(r'<br\s*/?>', '\n', md)
        md = re.sub(r'<[^>]+>', '', md)
        return md
    
    def _create_docx(self, text: str, path: Path) -> ConversionResult:
        try:
            from docx import Document
            doc = Document()
            for paragraph in text.split("\n"):
                doc.add_paragraph(paragraph)
            doc.save(path)
            return ConversionResult(success=True, output_path=str(path))
        except ImportError:
            return ConversionResult(
                success=False,
                error="python-docx required. Install: pip install universal-converter[documents]"
            )
    
    def _create_pdf(self, text: str, path: Path) -> ConversionResult:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import simpleSplit
            
            c = canvas.Canvas(str(path), pagesize=letter)
            width, height = letter
            y = height - 50
            
            for line in text.split("\n"):
                if y < 50:
                    c.showPage()
                    y = height - 50
                lines = simpleSplit(line, "Helvetica", 12, width - 100)
                for l in lines:
                    c.drawString(50, y, l)
                    y -= 15
                y -= 5
            c.save()
            return ConversionResult(success=True, output_path=str(path))
        except ImportError:
            return ConversionResult(
                success=False,
                error="reportlab required. Install: pip install universal-converter[documents]"
            )
    
    def convert_format(self, path: str, to_format: str, output: Optional[str] = None) -> ConversionResult:
        """Convert document to a specific format"""
        task = ConversionTask(
            source_path=Path(path),
            target_path=Path(output or Path(path).with_suffix(f".{to_format}")),
            source_format=Path(path).suffix[1:],
            target_format=to_format
        )
        return self.convert(task)
